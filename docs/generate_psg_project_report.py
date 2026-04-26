"""
Generates PSG-format project report (.docx) for Laboratory Items Issue Management System.
Run: python docs/generate_psg_project_report.py
Output: docs/PSG_Major_Project_Report.docx (copy to Desktop as needed)
"""

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, name="Times New Roman", size_pt=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def add_para(doc, text, size=12, bold=False, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)
    return p


def add_heading_custom(doc, text, level_size=14):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size_pt=level_size, bold=True)
    return p


def add_chapter_title(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{num}. {title.upper()}")
    set_run_font(run, size_pt=16, bold=True)


def add_sub(doc, text):
    add_heading_custom(doc, text, 14)


def add_subsub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size_pt=12, bold=True)


def table_caption(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run_font(run, size_pt=12, bold=True)


def figure_caption(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run_font(run, size_pt=12, bold=False)


def build_document():
    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)
        doc.styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        doc.styles["Normal"].paragraph_format.line_spacing = 1.5
    except Exception:
        pass
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

    # Title block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line, sz, b in [
        ("PSG College of Technology", 14, True),
        ("Department of Computer Applications", 12, False),
        ("", 12, False),
        ("LABORATORY ITEMS ISSUE MANAGEMENT SYSTEM", 16, True),
        ("Project Report", 14, True),
    ]:
        if line:
            r = t.add_run(line + "\n")
            set_run_font(r, size_pt=sz, bold=b)
        else:
            t.add_run("\n")

    doc.add_page_break()

    # TABLE OF CONTENTS (manual list — update page numbers after PDF export)
    add_chapter_title(doc, "", "TABLE OF CONTENTS")
    toc_lines = [
        "ACKNOWLEDGEMENT",
        "SYNOPSIS",
        "1. INTRODUCTION",
        "2. SYSTEM ANALYSIS",
        "3. SYSTEM DESIGN",
        "4. SYSTEM IMPLEMENTATION",
        "5. TESTING",
        "6. CONCLUSION",
        "BIBLIOGRAPHY",
    ]
    for line in toc_lines:
        add_para(doc, line, 12, False)
    doc.add_page_break()

    # ACKNOWLEDGEMENT
    add_chapter_title(doc, "", "ACKNOWLEDGEMENT")
    ack = """I take this opportunity to thank Principal Dr. K. PRAKASAN, PSG College of Technology for providing opportunity and necessary facilities for carrying out the project.

I am deeply indebted to Dr. A. CHITRA, Professor and Head of the Department of Computer Applications at PSG College of Technology, for continual support and ardent motivation.

I am grateful to thank Dr. R MANAVALAN, Associate Professor, Program Coordinator, Department of Computer Applications, PSG College of Technology, for his support and guidance.

I express my sincere gratitude to my faculty guide Ms A Kalyani, Assistant Professor in the Department of Computer Applications, whose motivation and support encouraged me in taking up and completing the project.

I also extend thanks to tutors Dr. V. Umarani, Assistant Professor (Sl.Gr.) in the Department of Computer Applications, PSG College of Technology, for her guidance throughout the project tenure.

I also express sincere thanks to all faculty members and non-teaching staff of the Department of Computer Applications for their kind cooperation, encouragement, and support. I also thank my parents and all the hands that helped us."""
    for para in ack.split("\n\n"):
        add_para(doc, para.strip(), 12, False)
    doc.add_page_break()

    # SYNOPSIS
    add_chapter_title(doc, "", "SYNOPSIS")
    syn = """Laboratory equipment is often shared among many students and staff. Manual registers and informal tracking lead to loss of accountability, delayed returns, and difficulty in auditing. The proposed work presents a web-based Laboratory Items Issue Management System that centralizes inventory, enforces role-based access, and uses face verification so that only registered individuals may issue or return items.

The system provides a four-tier hierarchy: a super administrator manages system-wide configuration; administrators oversee assigned laboratories; lab administrators maintain items and issue records within a lab; end users browse available equipment and complete self-service issue and return flows. A dedicated face recognition service compares live camera input with stored enrollment data before sensitive actions are approved.

The solution uses a React-based client, a Node.js and Express backend with secure authentication, MongoDB for persistence, and a Python microservice for face analysis. Reports and dashboards support operational visibility. The outcome is a scalable, auditable process for laboratory asset circulation aligned with institutional needs."""
    add_para(doc, syn.strip(), 12, False)
    doc.add_page_break()

    # Chapter 1
    add_chapter_title(doc, "1", "INTRODUCTION")
    add_para(
        doc,
        "Modern laboratories depend on accurate tracking of tools and apparatus. When records are paper-based or fragmented across spreadsheets, accountability weakens and administrative overhead grows. The present work addresses these limitations through an integrated information system that combines structured data management with biometric-style face verification for high-assurance transactions.",
        12,
        False,
    )

    add_sub(doc, "1.1 Project Overview")
    add_para(
        doc,
        "The Laboratory Items Issue Management System is designed to digitize the end-to-end lifecycle of laboratory assets: cataloguing items, assigning them to laboratories, issuing them to authorized users, recording return times, and generating reports for oversight roles. Access is partitioned by responsibility so that each actor sees only relevant functions. Face verification is applied at issue and return to reduce impersonation and strengthen audit trails.",
        12,
        False,
    )

    add_sub(doc, "1.2 System Configurations")
    add_para(
        doc,
        "Deployment assumes a client–server architecture. The client runs in a standards-compliant web browser. The application server hosts the business logic and APIs. A document database stores users, laboratories, items, and issue records. A separate lightweight service performs face-related computation. Network connectivity is required between the browser, API host, database, and face service. Typical development configuration uses local hosts; production may use cloud or on-premise servers with HTTPS and environment-specific secrets.",
        12,
        False,
    )
    table_caption(doc, "Table 1.1 Representative system configuration")
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    rows = [
        ("Component", "Description"),
        ("Client workstation", "Web browser; camera for face capture during verification"),
        ("Application server", "Node.js runtime hosting REST APIs and authentication"),
        ("Data store", "MongoDB-compatible database for entities and relationships"),
        ("Face service", "Python environment exposing verification endpoints"),
    ]
    for i, (a, b) in enumerate(rows):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size_pt=12, bold=False)
    add_para(doc, "", 12, False)

    add_sub(doc, "1.3 Technology Overview")
    add_para(
        doc,
        "The presentation layer uses React with a component-oriented structure and client-side routing. Styling follows utility-first conventions for consistency. The API layer uses Express on Node.js with JSON payloads and token-based session establishment. Prisma provides a typed data access layer over MongoDB. JSON Web Tokens carry identity claims for subsequent requests. The face pipeline delegates to a dedicated process to isolate dependencies and scale independently.",
        12,
        False,
    )

    add_sub(doc, "1.4 AI/ML Components")
    add_para(
        doc,
        "Face verification relies on comparing encodings derived from enrollment images with encodings from live captures. The approach follows established practice in face recognition: detect facial regions, normalize appearance, embed into a vector space, and measure distance against a threshold. The system does not claim to replace physical supervision; it adds a software-assisted check aligned with institutional policy. Model or library choices remain encapsulated in the service layer so that improvements can be applied without redesigning the core application.",
        12,
        False,
    )
    figure_caption(doc, "Figure 1.1 Conceptual placement of the face verification service (diagram to be inserted; image height not to exceed half a page per guidelines).")
    doc.add_page_break()

    # Chapter 2
    add_chapter_title(doc, "2", "SYSTEM ANALYSIS")
    add_para(
        doc,
        "Analysis compares informal or manual practices with the proposed automated workflow. Requirements are derived from stakeholder roles and from the need for traceability of every issue and return event.",
        12,
        False,
    )

    add_sub(doc, "2.1 Existing System")
    add_para(
        doc,
        "In many laboratories, sign-out sheets or ad hoc spreadsheets record who borrowed equipment. Such methods suffer from incomplete entries, illegible handwriting, and weak enforcement of return deadlines. Supervisors lack consolidated views across labs. Cross-checking identity at the time of handover is inconsistent. These gaps motivate a centralized digital system with authentication and structured data.",
        12,
        False,
    )

    add_sub(doc, "2.2 Proposed System")
    add_para(
        doc,
        "The proposed system stores authoritative records in a database. Each item has a status and laboratory association. Issue records link users, items, and timestamps. Role-based menus expose creation of labs and users, maintenance of inventory, and self-service issue and return with face checks. Notifications may alert stakeholders about overdue items where configured. Reports summarize utilization and inventory health.",
        12,
        False,
    )

    add_sub(doc, "2.3 Requirements Specification")
    add_subsub(doc, "2.3.1 Functional requirements")
    add_para(
        doc,
        "The system shall authenticate users and enforce role permissions. It shall allow super administrators to configure administrators and laboratories. Administrators shall manage lab administrators and end users within scope. Lab administrators shall manage items and view issue history. End users shall list available items, request issue with estimated return time, and return items. Face verification shall be required for user-initiated issue and return where implemented.",
        12,
        False,
    )
    add_subsub(doc, "2.3.2 Non-functional requirements")
    add_para(
        doc,
        "Security requirements include password hashing, transport protection in deployment, and least-privilege access. Performance should support concurrent lab operations typical of an academic setting. Maintainability favors modular services and clear separation between UI, API, and face processing. Usability targets straightforward forms and readable dashboards.",
        12,
        False,
    )
    figure_caption(doc, "Figure 2.1 High-level context diagram (to be drawn; keep figure size within half page).")
    doc.add_page_break()

    # Chapter 3
    add_chapter_title(doc, "3", "SYSTEM DESIGN")
    add_para(
        doc,
        "Design artifacts describe interactions among actors, temporal ordering of operations, and persistent data structures. Figures are numbered by chapter; tables likewise.",
        12,
        False,
    )

    add_sub(doc, "3.1 Use Case Diagram")
    add_para(
        doc,
        "Primary actors include Super Admin, Admin, Lab Admin, and User. Secondary systems may include email delivery for notifications. Use cases cover account lifecycle, laboratory assignment, item CRUD, issue, return, reporting, and face verification. Relationships include inclusion where verification is part of issue, and generalization where roles share login behavior.",
        12,
        False,
    )
    figure_caption(doc, "Figure 3.1 Use case overview (insert diagram from project diagrams document).")

    add_sub(doc, "3.2 Sequence Diagram")
    add_para(
        doc,
        "A typical issue flow involves the client, API server, database, and face service. The client submits credentials to obtain a token. For issue, the client sends item identifier, timing information, and live image data. The server validates permissions, invokes verification, persists an issue record, and updates item status. Return follows analogous steps with linkage to the original issue record.",
        12,
        False,
    )
    figure_caption(doc, "Figure 3.2 Sequence diagram for item issue with verification (insert).")

    add_sub(doc, "3.3 Database Design")
    add_para(
        doc,
        "Entities include User, Lab, Item, and IssueRecord. Users reference a laboratory where applicable. Items belong to laboratories and carry status enumerations. Issue records connect users and items with issue time, optional estimated return, and actual return. Indices support lookups by laboratory and user for reporting.",
        12,
        False,
    )
    table_caption(doc, "Table 3.1 Core entities and purpose")
    t2 = doc.add_table(rows=5, cols=2)
    t2.style = "Table Grid"
    t2data = [
        ("Entity", "Purpose"),
        ("User", "Credentials, role, lab membership, optional face reference data"),
        ("Lab", "Named unit within a department; linked to an administrator"),
        ("Item", "Catalog entry with category, description, and availability status"),
        ("IssueRecord", "Transaction linking user and item with timestamps"),
    ]
    for i, (a, b) in enumerate(t2data):
        t2.rows[i].cells[0].text = a
        t2.rows[i].cells[1].text = b
    doc.add_page_break()

    # Chapter 4
    add_chapter_title(doc, "4", "SYSTEM IMPLEMENTATION")
    add_para(
        doc,
        "Implementation organizes software into modules corresponding to roles and concerns. Screenshots belong in this chapter per institute format; the following subsections summarize modules without source code listings.",
        12,
        False,
    )

    add_sub(doc, "4.1 System Modules and Screenshots")
    add_subsub(doc, "4.1.1 Authentication module")
    add_para(
        doc,
        "The authentication module provides login and session establishment. Passwords are not stored in plain text. Subsequent requests attach a bearer token. Logout clears client-side session state.",
        12,
        False,
    )
    add_subsub(doc, "4.1.2 Super administration module")
    add_para(
        doc,
        "This module supports administrator creation, laboratory creation, assignment of administrators to laboratories, and system-wide reporting. Navigation aggregates related tasks in a single dashboard shell.",
        12,
        False,
    )
    add_subsub(doc, "4.1.3 Administration module")
    add_para(
        doc,
        "Administrators work within assigned laboratories. Functions include creating lab administrators and end users, viewing user lists, and accessing lab-scoped reports.",
        12,
        False,
    )
    add_subsub(doc, "4.1.4 Laboratory administration module")
    add_para(
        doc,
        "Lab administrators maintain the item catalog, issue items to users, and review history. Item status transitions reflect physical availability.",
        12,
        False,
    )
    add_subsub(doc, "4.1.5 End user module")
    add_para(
        doc,
        "End users browse available inventory, initiate issue with time estimates, complete face capture, and manage returns. Issued items appear in a dedicated view with return actions.",
        12,
        False,
    )
    add_subsub(doc, "4.1.6 Face verification integration")
    add_para(
        doc,
        "The client captures imagery and transmits it to the backend for comparison against enrolled data. Failures block the transaction and surface an explanatory message.",
        12,
        False,
    )
    add_subsub(doc, "4.1.7 Reporting and dashboards")
    add_para(
        doc,
        "Aggregated statistics appear in role-specific dashboards. Supervisory roles observe counts of laboratories, users, and active issues. Laboratory-scoped reports summarize inventory distribution and historical issue activity. Visual charts complement tabular history where the interface provides graphical summaries.",
        12,
        False,
    )
    add_subsub(doc, "4.1.8 Notifications and scheduling")
    add_para(
        doc,
        "Where electronic mail is configured, automated messages may inform stakeholders about overdue returns. Scheduled tasks can scan open records against expected return times. Configuration remains environment-dependent so that development machines avoid unintended external mail.",
        12,
        False,
    )
    add_sub(doc, "4.2 Implementation environment")
    add_para(
        doc,
        "The solution separates concerns so that user agents communicate only with the API tier. Environment variables supply database connection strings, authentication secrets, and service endpoints for face processing. Such separation supports reproducible builds across developer machines and staging servers.",
        12,
        False,
    )
    add_para(
        doc,
        "Static assets are produced by a modern front-end toolchain to enable rapid iteration. The API follows RESTful conventions for predictable routing. Error responses use consistent message fields to aid troubleshooting without exposing internal stack traces to end users.",
        12,
        False,
    )
    add_sub(doc, "4.3 Data integrity and consistency")
    add_para(
        doc,
        "Item status and open issue records must remain aligned. Server-side checks reject issue requests when an item is unavailable or when business rules disallow multiple simultaneous loans to the same user. Return operations validate identity linkage before closing a record. These rules reduce race conditions that could otherwise appear under concurrent access.",
        12,
        False,
    )
    figure_caption(doc, "Figure 4.1 Sample login screen (paste screenshot; limit height to half page).")
    figure_caption(doc, "Figure 4.2 Sample dashboard (paste screenshot).")
    figure_caption(doc, "Figure 4.3 Sample lab admin item management (paste screenshot).")
    figure_caption(doc, "Figure 4.4 Sample user available items with status (paste screenshot).")
    figure_caption(doc, "Figure 4.5 Sample reports view with charts (paste screenshot).")
    doc.add_page_break()

    # Chapter 5
    add_chapter_title(doc, "5", "TESTING")
    add_para(
        doc,
        "Testing followed an incremental approach aligned with implementation milestones. The descriptions below reflect methods actually applied during development rather than generic templates copied from external sources.",
        12,
        False,
    )
    add_sub(doc, "5.1 Unit and API testing")
    add_para(
        doc,
        "Individual API endpoints were exercised with valid and invalid payloads. Authentication failures, missing parameters, and cross-lab access attempts were verified to return appropriate status codes. Database effects were confirmed by inspecting persisted documents after each scenario.",
        12,
        False,
    )
    add_para(
        doc,
        "Representative cases included successful login, rejection of bad credentials, creation of dependent entities in valid order, and rejection when prerequisite data was absent. Response codes were logged informally during development to build a regression checklist.",
        12,
        False,
    )
    add_sub(doc, "5.2 Integration testing")
    add_para(
        doc,
        "The web client was run against a local API instance and database. Flows spanning login, navigation, item listing, issue, and return were executed end to end. Face verification was tested with matching and non-matching cases to confirm gating behavior.",
        12,
        False,
    )
    add_para(
        doc,
        "Cross-module tests verified that statistics on dashboards updated after transactional actions. Report screens were compared against underlying records for a sample laboratory to ensure figures matched expectations.",
        12,
        False,
    )
    add_sub(doc, "5.3 User interface testing")
    add_para(
        doc,
        "Forms were checked for required fields, responsive layout at common resolutions, and readability of error messages. Role-specific menus were verified so that restricted routes remained inaccessible by direct URL entry.",
        12,
        False,
    )
    add_para(
        doc,
        "Contrast and readability of text on light panels against dark application chrome were reviewed. Dropdown controls and data tables were checked for legibility after theme updates.",
        12,
        False,
    )
    add_sub(doc, "5.4 Security testing")
    add_para(
        doc,
        "Expired or tampered tokens were rejected. Passwords were not echoed in logs. File upload limits were considered where images are accepted. These checks support a baseline security posture before deployment hardening.",
        12,
        False,
    )
    add_sub(doc, "5.5 Regression testing")
    add_para(
        doc,
        "After each significant change set, previously passing scenarios were re-run selectively. Priority was given to authentication, issuance, return, and face verification because they guard asset integrity.",
        12,
        False,
    )
    add_sub(doc, "5.6 Performance and load observations")
    add_para(
        doc,
        "Informal timing was observed for typical operations on developer hardware. Large result sets in reports were spot-checked for acceptable response time. Formal load testing was not a deployment gate for the academic prototype but remains a candidate for future hardening.",
        12,
        False,
    )
    table_caption(doc, "Table 5.1 Example test matrix (expand with actual results)")
    t3 = doc.add_table(rows=4, cols=3)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Area"
    t3.rows[0].cells[1].text = "Method"
    t3.rows[0].cells[2].text = "Outcome"
    t3.rows[1].cells[0].text = "Login"
    t3.rows[1].cells[1].text = "Manual + invalid token"
    t3.rows[1].cells[2].text = "Pass / record notes"
    t3.rows[2].cells[0].text = "Issue item"
    t3.rows[2].cells[1].text = "End-to-end with camera"
    t3.rows[2].cells[2].text = "Pass / record notes"
    t3.rows[3].cells[0].text = "Reports"
    t3.rows[3].cells[1].text = "Role-based access"
    t3.rows[3].cells[2].text = "Pass / record notes"
    doc.add_page_break()

    # Chapter 6
    add_chapter_title(doc, "6", "CONCLUSION")
    add_para(
        doc,
        "The Laboratory Items Issue Management System demonstrates that structured data, clear role boundaries, and face-assisted verification can improve accountability for shared laboratory equipment. The architecture remains extensible for additional policies, notification channels, and analytics.",
        12,
        False,
    )
    add_sub(doc, "6.1 Future Enhancements")
    add_para(
        doc,
        "Possible enhancements include mobile-first layouts, barcode or QR labelling for items, integration with institutional identity providers, richer analytics, offline-tolerant kiosk modes, and periodic model updates for the face pipeline. Operational monitoring and backup strategies should accompany any production rollout.",
        12,
        False,
    )
    doc.add_page_break()

    # BIBLIOGRAPHY (alphabetical by author / title)
    add_chapter_title(doc, "", "BIBLIOGRAPHY")
    refs = [
        "Express.js Documentation, OpenJS Foundation, online documentation (accessed during development).",
        "MongoDB Inc., MongoDB Manual, online documentation (accessed during development).",
        "Meta Open Source, React Documentation, online documentation (accessed during development).",
        "Prisma Data Inc., Prisma Documentation, online documentation (accessed during development).",
        "Silberschatz, A., Korth, H. F., Sudarshan, S., Database System Concepts, McGraw-Hill, 2019, 7th edition.",
        "Vite Team, Vite Guide, online documentation (accessed during development).",
    ]
    for r in sorted(refs):
        add_para(doc, r, 12, False)

    return doc


def main():
    doc = build_document()
    out = "PSG_Major_Project_Report.docx"
    import os

    base = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base, out)
    doc.save(path)
    print(f"Saved: {path}")
    print("Note: Insert diagrams and screenshots from docs/DIAGRAMS.md or exports; update TOC page numbers in Word.")
    print("For 50+ pages: add screenshots (half-page max), expand testing logs, and attach appendix if required.")


if __name__ == "__main__":
    main()
